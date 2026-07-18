"""Course material ingestion, chunking and persistent FAISS retrieval."""

from __future__ import annotations

import hashlib
import json
import pickle
import re
import shutil
import threading
from pathlib import Path
from typing import Iterable

from langchain_community.vectorstores import FAISS
from langchain_community.vectorstores.utils import DistanceStrategy
from langchain_core.documents import Document
from langchain_core.embeddings import Embeddings

from config import SUPPORTED_EXTENSIONS, Settings
from models import RetrievedChunk


class RAGService:
    """Owns the local material library and vector index."""

    def __init__(self, settings: Settings, embeddings: Embeddings | None = None):
        self.settings = settings
        self.settings.ensure_directories()
        self._embeddings = embeddings
        self._vector_store: FAISS | None = None
        self._index_stale = False
        self._lock = threading.RLock()
        self.manifest_path = self.settings.data_dir / "index_manifest.json"
        self._load_index()

    @property
    def embeddings(self) -> Embeddings:
        if self._embeddings is None:
            import httpx
            from langchain_openai import OpenAIEmbeddings

            self.settings.validate_embedding_api()
            embedding_kwargs = {}
            if self.settings.embedding_dimensions is not None:
                embedding_kwargs["dimensions"] = self.settings.embedding_dimensions
            http_client = httpx.Client(verify=self.settings.verify_ssl, timeout=120)
            self._embeddings = OpenAIEmbeddings(
                api_key=self.settings.embedding_api_key,
                base_url=self.settings.embedding_client_base_url,
                model=self.settings.embedding_model,
                timeout=120,
                max_retries=2,
                http_client=http_client,
                tiktoken_enabled=False,
                check_embedding_ctx_length=False,
                **embedding_kwargs,
            )
        return self._embeddings

    def _embedding_signature(self) -> dict:
        return {
            "provider": "openai-compatible",
            "base_url": self.settings.embedding_base_url,
            "model": self.settings.embedding_model,
            "dimensions": self.settings.embedding_dimensions,
        }

    def _index_matches_embeddings(self) -> bool:
        return self._read_manifest().get("embedding") == self._embedding_signature()

    def _load_index(self) -> None:
        index_file = self.settings.vector_dir / "index.faiss"
        metadata_file = self.settings.vector_dir / "index.pkl"
        material_paths = self._material_paths()
        if not material_paths:
            # The material directory is the source of truth.  In particular,
            # do not expose or query a vector index left by an earlier run.
            self._vector_store = None
            self._index_stale = False
            return
        if not self._manifest_matches_materials():
            self._vector_store = None
            self._index_stale = True
            return
        if index_file.exists() and metadata_file.exists():
            if not self._index_matches_embeddings():
                self._vector_store = None
                self._index_stale = True
                return
            # faiss.write_index cannot open Unicode paths on some Windows builds.
            # Byte serialization lets pathlib handle the Chinese workspace path.
            import faiss
            import numpy as np

            buffer = np.frombuffer(index_file.read_bytes(), dtype=np.uint8)
            index = faiss.deserialize_index(buffer)
            with metadata_file.open("rb") as stream:
                docstore, index_to_docstore_id = pickle.load(stream)
            self._vector_store = FAISS(
                self.embeddings,
                index,
                docstore,
                index_to_docstore_id,
                normalize_L2=True,
                distance_strategy=DistanceStrategy.EUCLIDEAN_DISTANCE,
            )
        else:
            self._index_stale = True

    def _save_index(self) -> None:
        if self._vector_store is None:
            return
        import faiss

        index_bytes = faiss.serialize_index(self._vector_store.index).tobytes()
        (self.settings.vector_dir / "index.faiss").write_bytes(index_bytes)
        with (self.settings.vector_dir / "index.pkl").open("wb") as stream:
            pickle.dump(
                (
                    self._vector_store.docstore,
                    self._vector_store.index_to_docstore_id,
                ),
                stream,
            )

    @staticmethod
    def _sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as stream:
            for block in iter(lambda: stream.read(1024 * 1024), b""):
                digest.update(block)
        return digest.hexdigest()

    def _read_manifest(self) -> dict:
        if not self.manifest_path.exists():
            return {"files": []}
        try:
            return json.loads(self.manifest_path.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            return {"files": []}

    def _material_paths(self) -> list[Path]:
        return [
            path
            for path in sorted(self.settings.materials_dir.iterdir())
            if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
        ]

    def _manifest_matches_materials(self) -> bool:
        records = self._read_manifest().get("files", [])
        records_by_name = {
            item.get("name"): item for item in records if isinstance(item, dict)
        }
        paths = self._material_paths()
        if set(records_by_name) != {path.name for path in paths}:
            return False
        return all(
            records_by_name[path.name].get("sha256") == self._sha256(path)
            for path in paths
        )

    def list_materials(self) -> list[dict]:
        records = self._read_manifest().get("files", [])
        records_by_name = {
            item.get("name"): item for item in records if isinstance(item, dict)
        }
        result = []
        for path in self._material_paths():
            file_hash = self._sha256(path)
            record = records_by_name.get(path.name)
            if record and record.get("sha256") == file_hash:
                result.append(record)
            else:
                result.append(
                    {
                        "name": path.name,
                        "sha256": file_hash,
                        "chunks": 0,
                        "size": path.stat().st_size,
                    }
                )
        return result

    def add_materials(self, paths: Iterable[str | Path]) -> dict:
        added, skipped, errors = [], [], []
        existing_hashes = {item.get("sha256") for item in self.list_materials()}
        for raw_path in paths:
            source = Path(raw_path).expanduser().resolve()
            if not source.exists() or source.suffix.lower() not in SUPPORTED_EXTENSIONS:
                errors.append(f"不支持或不存在：{source}")
                continue
            file_hash = self._sha256(source)
            if file_hash in existing_hashes:
                skipped.append(source.name)
                continue
            target = self.settings.materials_dir / source.name
            if target.exists() and self._sha256(target) != file_hash:
                target = target.with_name(f"{target.stem}_{file_hash[:8]}{target.suffix}")
            if source != target:
                shutil.copy2(source, target)
            added.append(target.name)
            existing_hashes.add(file_hash)
        if added or self._index_stale:
            try:
                self.rebuild_index()
            except Exception as exc:
                errors.append(f"索引构建失败：{exc}")
        return {"added": added, "skipped": skipped, "errors": errors}

    def remove_material(self, filename: str) -> bool:
        target = self.settings.materials_dir / Path(filename).name
        if not target.exists():
            return False
        target.unlink()
        self.rebuild_index()
        return True

    def rebuild_index(self) -> dict:
        with self._lock:
            all_chunks: list[Document] = []
            file_records: list[dict] = []
            for path in self._material_paths():
                documents = self._load_file(path)
                chunks = self._split_documents(documents)
                all_chunks.extend(chunks)
                file_records.append(
                    {
                        "name": path.name,
                        "sha256": self._sha256(path),
                        "chunks": len(chunks),
                        "size": path.stat().st_size,
                    }
                )
            self._vector_store = None
            if all_chunks:
                self._vector_store = FAISS.from_documents(
                    all_chunks,
                    self.embeddings,
                    distance_strategy=DistanceStrategy.EUCLIDEAN_DISTANCE,
                    normalize_L2=True,
                )
                self._save_index()
            else:
                for file in self.settings.vector_dir.glob("*"):
                    if file.is_file():
                        file.unlink()
            manifest = {
                "embedding": self._embedding_signature(),
                "files": file_records,
                "total_chunks": len(all_chunks),
            }
            self.manifest_path.write_text(
                json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8"
            )
            self._index_stale = False
            return manifest

    def _load_file(self, path: Path) -> list[Document]:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            # The aggregate document_loaders package imports optional ML stacks
            # on Windows, so defer it until a PDF is actually ingested.
            from langchain_community.document_loaders.pdf import PyPDFLoader

            pages = PyPDFLoader(str(path)).load()
            for page in pages:
                page.metadata.update(
                    {
                        "source": path.name,
                        "source_path": str(path),
                        "chapter": self._guess_heading(page.page_content)
                        or f"第 {page.metadata.get('page', 0) + 1} 页",
                    }
                )
            return pages
        if suffix == ".docx":
            return self._load_docx(path)
        text = self._read_text(path)
        return self._section_documents(text, path)

    @staticmethod
    def _read_text(path: Path) -> str:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return path.read_text(encoding=encoding)
            except UnicodeDecodeError:
                continue
        raise UnicodeError(f"无法识别文件编码：{path.name}")

    def _load_docx(self, path: Path) -> list[Document]:
        try:
            from docx import Document as DocxDocument
        except ImportError as exc:
            raise RuntimeError("读取 DOCX 需要安装 python-docx") from exc
        docx = DocxDocument(str(path))
        sections: list[Document] = []
        heading, buffer = "文档开始", []
        for paragraph in docx.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                if buffer:
                    sections.append(self._make_document("\n".join(buffer), path, heading))
                heading, buffer = text, []
            else:
                buffer.append(text)
        if buffer:
            sections.append(self._make_document("\n".join(buffer), path, heading))
        return sections

    def _section_documents(self, text: str, path: Path) -> list[Document]:
        heading_pattern = re.compile(r"^(#{1,6}\s+.+|第[一二三四五六七八九十\d]+[章节].*)$")
        heading, buffer, sections = "文档开始", [], []
        for line in text.splitlines():
            clean = line.strip()
            if heading_pattern.match(clean):
                if buffer:
                    sections.append(self._make_document("\n".join(buffer), path, heading))
                heading, buffer = clean.lstrip("# "), []
            else:
                buffer.append(line)
        if buffer:
            sections.append(self._make_document("\n".join(buffer), path, heading))
        return [doc for doc in sections if doc.page_content.strip()]

    @staticmethod
    def _make_document(text: str, path: Path, heading: str) -> Document:
        return Document(
            page_content=text.strip(),
            metadata={
                "source": path.name,
                "source_path": str(path),
                "chapter": heading,
            },
        )

    @staticmethod
    def _guess_heading(text: str) -> str:
        for line in text.splitlines()[:12]:
            line = line.strip()
            if 2 <= len(line) <= 60 and (
                line.startswith("第") or re.match(r"^\d+(\.\d+)*[、.\s]", line)
            ):
                return line
        return ""

    def _split_documents(self, documents: list[Document]) -> list[Document]:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.chunk_size,
            chunk_overlap=self.settings.chunk_overlap,
            separators=["\n## ", "\n### ", "\n", "。", "；", "，", " "],
        )
        chunks = splitter.split_documents(documents)
        for index, chunk in enumerate(chunks):
            source = chunk.metadata.get("source", "unknown")
            chunk.metadata["chunk_id"] = f"{source}:{index + 1}"
        return chunks

    def search(
        self,
        query: str,
        top_k: int | None = None,
        source: str | None = None,
        chapter: str | None = None,
    ) -> list[RetrievedChunk]:
        if self._vector_store is None and self._index_stale and self.list_materials():
            self.rebuild_index()
        if not self._vector_store or not query.strip():
            return []
        wanted = top_k or self.settings.top_k
        # Embeddings are unit-normalized, therefore cosine similarity can be
        # recovered from squared L2 distance: cos(a, b) = 1 - d² / 2.
        raw = self._vector_store.similarity_search_with_score(query, k=max(wanted * 4, 12))
        results: list[RetrievedChunk] = []
        for doc, raw_score in raw:
            metadata = doc.metadata
            if source and metadata.get("source") != source:
                continue
            if chapter and chapter not in metadata.get("chapter", ""):
                continue
            page = metadata.get("page")
            location = (
                f"第 {int(page) + 1} 页"
                if page is not None
                else metadata.get("chapter", "未标注章节")
            )
            score = max(0.0, min(1.0, 1.0 - float(raw_score) / 2.0))
            results.append(
                RetrievedChunk(
                    content=doc.page_content,
                    source=metadata.get("source", "未知资料"),
                    location=location,
                    chapter=metadata.get("chapter", ""),
                    chunk_id=metadata.get("chunk_id", ""),
                    score=score,
                )
            )
            if len(results) >= wanted:
                break
        return results

    def list_chapters(self, source: str | None = None) -> list[str]:
        if self._vector_store is None and self._index_stale and self.list_materials():
            self.rebuild_index()
        if not self._vector_store:
            return []
        docstore = getattr(self._vector_store, "docstore", None)
        documents = getattr(docstore, "_dict", {}).values()
        chapters: list[str] = []
        seen: set[str] = set()
        for doc in documents:
            metadata = getattr(doc, "metadata", {})
            if source and metadata.get("source") != source:
                continue
            chapter = str(metadata.get("chapter", "")).strip()
            if (
                not chapter
                or chapter == "文档开始"
                or chapter in seen
                or not self._is_selectable_chapter(chapter)
            ):
                continue
            chapters.append(chapter)
            seen.add(chapter)
        return chapters

    @staticmethod
    def _is_selectable_chapter(chapter: str) -> bool:
        return bool(
            re.match(r"^第\s*[一二三四五六七八九十\d]+\s*[章节]", chapter)
            or re.match(r"^实验\s*[一二三四五六七八九十\d]+", chapter)
            or re.match(r"^\d+(\.\d+)*[、.\s]", chapter)
        )
